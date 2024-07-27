import streamlit as st
import PyPDF2 as pdf
import google.generativeai as genai
from docx import Document
from io import BytesIO


# ---------------------------------- Initializing ---------------------------------- #
# Access the API key from Streamlit secrets
api_key = st.secrets["General"]["GOOGLE_API_KEY"]
genai.configure(api_key=api_key)


# --------------------- Functions  --------------------- #

def local_css(file_name):
    with open(file_name) as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

# Load custom CSS
local_css("style.css")


# Gemini Pro Response
def get_gemini_response(input):
    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(input)
    # Extract the generated text from the response object
    if hasattr(response, 'candidates') and response.candidates:
        return response.candidates[0].content.parts[0].text
    else:
        return "No response generated"


def input_pdf_text(uploaded_file):
    reader = pdf.PdfReader(uploaded_file)
    text = ''
    for page in reader.pages:
        text += page.extract_text()
    return text


# Function to save response to a docx file
def save_response_to_docx(response_text, filename="cover_letter.docx"):
    doc = Document()
    doc.add_heading("Cover Letter", level=1)
    doc.add_paragraph(response_text)
    doc.save(filename)
    print(f"File {filename} saved successfully.")  # Debug statement
    

# --------------------- Side Bar --------------------- #

with st.sidebar:
    
    # Logo
    st.image("assests\done.png", width=135)
    
    # File upload for Resume
    uploaded_file = st.file_uploader("Upload Resume:", type='pdf', help='Please upload resume in PDF format')
    

    
    # Generate Cover Letter Button
    
    job = st.text_area("Paste the Job Description")
    generate_cover_letter = st.button("Generate Cover Letter")
    
    # Footer
    st.markdown("""---""")
  
    # LinkedIn and Twitter Icons
    linkedin_url = "https://www.linkedin.com/in/precious0728/"
    x_url = "https://twitter.com/presydon"
    youtube = "https://www.youtube.com/shorts/Amh7BPEUEXc"
    st.markdown(
        f'<div class="icon-row">'
        f'<a href="{linkedin_url}" target="_blank">'
        f'<img src="https://img.icons8.com/color/48/000000/linkedin.png" class="icon animated-icon"></a>'
        f'<a href="{x_url}" target="_blank">'
        f'<img src="https://img.icons8.com/color/48/000000/twitter.png" class="icon animated-icon"></a>'
        f'</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="sidebar-footer">Built by Precious Victor</div>', 
        unsafe_allow_html=True
        )


# --------------------- Custom CSS --------------------- #
css = """
<style>
.sidebar-text {
    font-size: 0.9em;
}
.sidebar-footer {
    font-size: 0.9em;
}

.sidebar-footer {
    font-size: 0.9em;
    text-align: center;
    display: flex;
    justify-content: center;
    align-items: center;
    margin-top: 20px;

.icon-row {
    display: flex;
    align-items: center;
}
.icon {
    width: 30px;
    height: 30px;
    transition: transform 0.3s;
}
.icon:hover {
    transform: scale(1.1);
}
</style>
"""

st.markdown(css, unsafe_allow_html=True)

# --------------------- Main Page --------------------- #

# First view
st.header("Resume Application Tracking System")
st.button('What is Resume ATS', on_click=lambda: st.write('An Applicant Tracking System (ATS) is software that automates the hiring process...'))

# Response section
jd = st.chat_input("Paste the Job Description")

# --------------------- Prompt --------------------- #
prompt_template = """
As a seasoned Application Tracking System (ATS) with comprehensive knowledge of various industries and career paths, you will:
1. Assess the candidate's resume {CV} against the provided Job Description {jd}, ensuring alignment with the role's requirements.
2. Provide actionable feedback to strengthen the resume, addressing crucial gaps and omissions in the competitive job market.
3. Extract and list the missing skills from the {jd} that are not present in the {CV}, enabling targeted improvements.
4. Provide a job description match percentage at the end of your review, give a proper grade not a range Example 61 and not 50-60
5. If Job description match is above 75%, suggest the user to CLICK- Generate CV
"""

prompt_2 = """"

You are a job applicant, based on this job description {jd} do the following
    1. identify the biggest challenge someone in this position would face day-to-day
    2. then review through the resume {cv} generate an attention grabbinh hook for the cover letter that highlights your experience and qualifications in a way that shows you empathize and can successgully take on the challenge of the job.
    3.  consider incorporating specific examplesof how you've tackled these challenges in your past work and explore creative ways to express your enthusiasm for the opportunity. keep your hook within 100 words.
with all this information, write a cover letter applying for the role, generate the cover letter based on the resume {cv} and hook created earlier, keep it concise, clear and professional, finally, within 250 words
"""

# --------------------- streamlit actions --------------------- #

if jd and uploaded_file:
    CV = input_pdf_text(uploaded_file)
    prompt = prompt_template.format(CV=CV, jd=jd)
    try:
        response = get_gemini_response(prompt)
        st.write(response)
    except Exception as e:
        st.error(f'An Error Occurred: {e}')

if generate_cover_letter:
    if job and uploaded_file:
        try:
            # Read PDF resume
            cv = input_pdf_text(uploaded_file)
            print(f"Extracted text from resume: {cv[:100]}...")  # Debug statement
            
            # Generate response using Google Generative AI
            prompt = prompt_2.format(jd=jd, cv=cv)
            model = genai.GenerativeModel('gemini-pro')
            response = model.generate_content(prompt)
            response_text = response.candidates[0].content.parts[0].text
            print(f"Generated response: {response_text[:100]}...")  # Debug statement
            
            # Save response to docx file
            save_response_to_docx(response_text)
            
            # Provide download link for the docx file
            with open("cover_letter.docx", "rb") as file:
                btn = st.download_button(label="Download Cover Letter", data=file, file_name="cover_letter.docx")
                print("Download button created.")  # Debug statement
        
        except Exception as e:
            st.error(f"An error occurred: {e}")
            print(f"An error occurred: {e}")  # Debug statement
    else:
        if not job:
            st.error("Please provide a job description.")
        if not uploaded_file:
            st.error("Please upload a resume file.")
